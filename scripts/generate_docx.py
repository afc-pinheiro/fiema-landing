import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set explicit margins (padding) for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    """Set background color (shading) for a cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_paragraph_runs(p, text):
    """Parse inline formatting like bold (**text**) and markdown links [text](url) and add as runs."""
    # Pattern to find bold tags **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            # Check for link inside bold
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', bold_text)
            if link_match:
                run = p.add_run(f"{link_match.group(1)} ({link_match.group(2)})")
                run.bold = True
            else:
                run = p.add_run(bold_text)
                run.bold = True
        else:
            # Check for links in regular text
            subparts = re.split(r'(\[.*?\]\(.*?\))', part)
            for subpart in subparts:
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', subpart)
                if link_match:
                    p.add_run(f"{link_match.group(1)} ({link_match.group(2)})")
                else:
                    p.add_run(subpart)

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles configuration
    # Normal / Body style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Check if we are parsing a table
        if line.strip().startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # We reached the end of the table
            process_table(doc, table_rows)
            table_rows = []
            in_table = False
            
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
            
        # Horizontal Rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('__________________________________________________________________')
            run.font.color.rgb = None # default gray
            i += 1
            continue
            
        # Title (Heading 1 / Document Title)
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = None # default dark color
            i += 1
            continue
            
        # Heading 2
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue
            
        # Heading 3
        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
            
        # Bullet list item
        if line.strip().startswith('* ') or line.strip().startswith('- '):
            clean_line = re.sub(r'^\s*[\*\-]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_runs(p, clean_line)
            i += 1
            continue
            
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_paragraph_runs(p, line.strip())
        i += 1
        
    # In case the file ended with a table
    if in_table and table_rows:
        process_table(doc, table_rows)
        
    doc.save(docx_path)
    print(f"Document saved to {docx_path}")

def process_table(doc, rows):
    """Convert parsed markdown table rows to a structured word table."""
    # Clean rows and split by pipe
    cleaned_rows = []
    for r in rows:
        cells = [c.strip() for c in r.split('|')]
        # Remove empty cells at start and end due to leading/trailing pipes
        if cells and not cells[0]:
            cells.pop(0)
        if cells and not cells[-1]:
            cells.pop()
        cleaned_rows.append(cells)
        
    if not cleaned_rows:
        return
        
    # Check if there is a divider row (like |---|---|) and remove it
    if len(cleaned_rows) > 1 and all(re.match(r'^:?-+:?$', c) for c in cleaned_rows[1]):
        cleaned_rows.pop(1)
        
    num_cols = len(cleaned_rows[0])
    num_rows = len(cleaned_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # Auto adjust column widths nicely
    # Let's set some proportional widths if there are 4 columns (like our Sprints table)
    # Col 1: Phase (1.2 in), Col 2: Período (0.8 in), Col 3: Sprints (0.8 in), Col 4: Detalhamento (3.7 in)
    col_widths = []
    if num_cols == 4:
        col_widths = [Inches(1.3), Inches(0.8), Inches(0.8), Inches(3.6)]
    
    for row_idx, row_data in enumerate(cleaned_rows):
        row = table.rows[row_idx]
        is_header = (row_idx == 0)
        
        # Set height or prevent row splitting
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if is_header:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for col_idx, text in enumerate(row_data):
            # Ensure cell exists (handle mismatched columns safely)
            if col_idx >= len(row.cells):
                continue
            cell = row.cells[col_idx]
            
            # Width adjustment
            if col_widths and col_idx < len(col_widths):
                cell.width = col_widths[col_idx]
                
            # Content
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            # Run
            add_paragraph_runs(p, text)
            
            # Format header row
            if is_header:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "1F497D") # Dark Blue Theme Accent
                # set text color to white
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_idx % 2 == 1:
                # Zebra striping for odd rows
                set_cell_shading(cell, "F2F5F8")
                
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

if __name__ == '__main__':
    md_path = 'docs/proposta_comercial_fiema.md'
    docx_path = 'proposta_comercial_fiema.docx'
    if os.path.exists(md_path):
        markdown_to_docx(md_path, docx_path)
    else:
        print(f"Error: {md_path} not found.")
